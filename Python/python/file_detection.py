"16/08/2025"
import os 
from tkinter import messagebox as message
from docx import Document #type: ignore
from docx.shared import Pt
import json 
from typing import Literal, Any,Iterable,overload
from watch import MyWatch #type: ignore
from my_translator_class import MYTranslator #type: ignore
from tkinter import Tk
import pandas as pd 
import traceback

"""
MODULE:
    This Module provided all the necesary to work with file in the current 
    computer. This module provided FileManager class. Where you can create, 
    read and write simple thing in file like: 
        - word files (.docx)
        -excel files (.xlsx)
        - json files (.json)
        -text files  (.txt)

    Also give provided a biniary search to look all the iterable in a effiency way

UPDATES:
    16/08/2025
    20/08/2025
    25/08/2025
    


"""







def my_binary(iterable:Iterable,
             target:Any)->int:
  """
  This is my own attempt of doing binary search. 
  \nDescription: 
        This will give us only the index of the target value, return -1 if the value does
        not exist or occurs a problem.
        To make this funtiob more optimized, you need to give it a sorted list already
    \nReturn: 
        the index of the target at hand as a interget, -1 otherwise

    \nParameters: 
        iterable: this is the iterable to follow through
        target: this is the value you want to know the index

    \nExample: 
    >>> my_binary([1,2,3,4,5],3)
        output: 2

  
  """ 
  start:int= 0
  end:int=len(iterable) -1 #type: ignore

  try: 

    while start <= end: 
      mid:int = (start + end)//2 

      if iterable[mid]== target: return mid #type: ignore

      elif iterable[mid]< target: start= mid +1 #type: ignore
        
      else: end= mid -1

  except: return -1

  return -1

def my_dict_binary(dictionaries:list[dict[str,Any]],
                   key:str,
                   target:str)->int: 
    """
    Performs a binary search on a sorted list of dictionaries.
    
    NOTE: The list of dictionaries MUST be sorted by the given 'key'
    before calling this function.
    
    Example: sorted_dicts.sort(key=lambda d: d['age'])
    """
    
 
    start:int=0
    end:int=len(dictionaries)-1

    while start <= end:
        mid:int= (start + end) // 2

        try: 

            middle_value = dictionaries[mid][key]

        except KeyError: 
            return -1
            
        if middle_value == target: return mid

        if middle_value < target: start = mid +1

        else: end = mid -1


    return -1

   

class FileManager():

    def __init__(self)->None: 
        self._with_traceback:bool=False
  

    #Commons methods to use everywhere
    @staticmethod
    def get_name(filePath)->str:
        """This will get the name of the file"""
        return os.path.basename(filePath)

    @staticmethod
    def path_desktop(fileName:str|None=None)->str:
        """This will only works with my pc, btw, dont forget to write a type (eg, .txt .docx)"""

        base_path = os.path.join(os.path.expanduser('~'), 'Desktop\\')

        if fileName is None:
            return base_path
        
        return os.path.join(base_path, fileName)

    @staticmethod
    def path_downloads(fileName:str|None=None)->str:
        if fileName is None: 
            return "C:\\Users\\Usuario Pc\\Downloads\\"

        return "C:\\Users\\Usuario Pc\\Downloads\\"  + fileName

    @staticmethod
    def path_lp(fileName:str|None=None)->str:
        if fileName is None: 
            return "C:\\Users\\Usuario Pc\\Desktop\\Learning_Python\\"

        return "C:\\Users\\Usuario Pc\\Desktop\\Learning_Python\\" + fileName


    @staticmethod
    def get_type(fileName:str):
        root, extension = os.path.splitext(fileName)
        return extension
  
    
    def translate_text(self,fileName:str,languages:list[str]=["english"]): 
        """
        This will try to get all the text from a text docuemnt and it will translat it 
        into one language in the list of languages. 

        Will create a new file name 'Traslated {name}' in the local folder
        
        """
        try: 
            text_=self.read_file(fileName)
            name:str=self.get_name(fileName)

            if not isinstance(text_,str): 
                print("The file has no text to translate")
                return 
        

            new_txt:str=MYTranslator().get_translation(text_,languages)

            self.create_file(f"Traslated {name}",new_txt)
            print("Done")

            self.storage_historial(fileName,"Translate")

        except Exception as e: 
            print("Problem at: "  + self.translate_text.__name__)
            print(e)


    def get_file_size(self,fileName)->int: 
        """
        This funtion will calculate the size of a file. To get the real size, later to 
        return the value you can divided by 100

        Return 
            The size of the file, or -1 if the file was not found
        
        """

        try: 
            size: int= os.path.getsize(fileName)
        except: 
            print(f"The File was {self.get_name(fileName)} not found")
            size:int =-1 #type: ignore
        finally:
            return size



    #General methods
    def read_file(self,fileName:str): 
        if fileName.endswith(".docx"):
            return self.read_word(fileName)
        
        if fileName.endswith(".txt"):
            return self.read_txt(fileName)
        
        if fileName.endswith(".json"):
            return self.read_json(fileName)
        
    def write_file(self,fileName:str,text:str): 
        if fileName.endswith(".docx"):
            return self.write_word(fileName,text)
        
        if fileName.endswith(".txt"):
            return self.write_txt(fileName,text)

    def create_file(self,fileName:str,text: str | None = None): 
        if fileName.endswith(".docx"):
            return self.create_word(fileName,text)
        
        if fileName.endswith(".txt"):
            return self.create_txt(fileName,text)
        

        


    #Method to storage all the info
    def storage_historial(self, fileName:str,reason:str="None")->None:
        """Here I just storage all file i open so far into a json file"""
        if not self.file_exist(fileName): self.create_json("Historial file.json")



        information:dict[str,Any]={
            "Day": MyWatch.get_local_time(),
            "Hour": MyWatch.get_local_hour(),
            "Reason": reason.capitalize(),
            "Path": fileName,
            "Type": self.get_type(self.get_name(fileName))
        }


        self.write_json("Historial file.json",information) 



    


    #Manage the files
    def file_exist(self,file_name:str)->bool:
        """This will be True if the file exist, False otherwise, so be sure to type it well
        (eg., Hello.txt)
        """
        return os.path.exists(file_name)
   
    def get_file_names(self,folderPath:str)->list[str]:
        """This will get a list of the names of every file in a specific folder"""
        if not self.file_exist(folderPath):
            message.showerror("Error","The folder does not exist, try another one")
            print( "Problem at: " +  self.get_file_names.__name__)
            return ["N/A"]

        

        file_names:list=[]

        for names in os.listdir(folderPath):
            file_=os.path.join(folderPath,names)
            if os.path.isfile(file_):
                file_names.append(names)
        
        file_names.sort()
        return file_names

    def open_file(self,fileName:str)->None: 

        try: 
            os.startfile(fileName)
            self.storage_historial(fileName,"Open")

        except Exception as e : 
            print( "Problem at: " +  self.open_file.__name__)
            message.showerror("Error",F"{e}")

    def search_file(self,name:str,open_:bool=True)->str:
        """
        This will attempt to  open  the target file, Using ONLY the file name
        through searching in all avaible path
        if the parameter open_ is True. However, if the file is found will return the path 
        of the file.

        Moreover, it will show a message error if the file was not found.

        return: 
            A string with the path of the file or 'N/A' if the file is not found
        
        """

        paths=[self.path_desktop(),self.path_downloads(),self.path_lp()]
        
        for folders in paths:

            files=self.get_file_names(folders)
            index=my_binary(files,name)
      

            if index != -1: 

                filename=folders + files[index]

                if open_: 
                    self.open_file(filename)
                    return  filename

                return filename
     

        message.showerror("File Dont found", f"The file '{name}' was not found anywhere")
        return "N/A"


    #Excel 
    def write_excel(self,
                            file_name:str,
                            sheet_name_:str='sheet1',
                            add_row:dict[Any,Any] |None=None,
                            with_index:bool=False)->bool:
        """
        DESCRIPTION:    
            This funtion will add new rows base of the info given, it will show a message if 
            a problem occurs or do nothing if without value is provied. 
        
        PARAMETERS:

            file_name: this only must be a string,could can chose whether to writer 
            .xlsx or not because  the .xlsx with be added automally
            sheet_name_: this only provided the sheet name, sheet1 otherwise
            with_index: if you want index or not, False as a default 
            add_row: this is the main info that you wanna add. it must be a dictionary
            with string as key, list as value

        
        RETURN: 
            An bool (True) if the Operation was succesfully Done. Otherewise means 
            the operation Fail
        
        """
        #1. write the documentations


        #2. do nothing if there is no data
        if not add_row:
            print("New rows was not provided")
            return False
        
        #3. this is a filter if the file name have xlsx in the end of the name, ro provide any erros 
        if not file_name.endswith(".xlsx"):
            file_name=file_name + ".xlsx"

        #this will calculate the size of the file 
        if self.get_file_size(file_name) <= 0: 
            return False


        

        try: 
            #4. if the file indeed exist 
            if os.path.exists(file_name):
                all_sheets: dict =pd.read_excel(file_name, sheet_name=None, engine='openpyxl')

                #Get the specif sheet we want, or it will create a new empty DateFrame object
                df:pd.DataFrame = all_sheets.get(sheet_name_,pd.DataFrame())
        

            #4.1. if the file does not exist, will create a new empty DateFrame object
            else: 
                all_sheets:dict={} #type: ignore
                df = pd.DataFrame()

            #5. create a new DateFrame with the data provide: add_row
            nw_row=pd.DataFrame(add_row)

            #6. this method will add a new ROW in the sheet at hand
            df = pd.concat([df,nw_row],ignore_index=True)

            #7.Here is update the value ( is the same like a normal dict)
            all_sheets[sheet_name_] = df
            
            #The main code of the function: 
            #this will create a new file if the file is not found or modify the file and the sheet
            #provided
            with pd.ExcelWriter(file_name,mode='w',engine='openpyxl') as writer:#type: ignore
                for name, dataframe in all_sheets.items():
                    dataframe.to_excel(writer, sheet_name=name, index=with_index)

            
                print( "Complet", "the modification was sucefully done!")
                self.storage_historial(file_name,"Write")

                return True

        #if a problem occurs, this except blocks will run 
        except PermissionError:
            print("Permission Error\n" + "Maybe you forgot to close the current file?")
            return False
        except Exception as e:
            print(f"Problem at:{self.write_excel.__name__}")
            print(e)
            if self._with_traceback:
                traceback.print_exc()
            return False
        
       




   #words files
    @staticmethod
    def word_styles(fileName:str)->list[str]:
        doc = Document(fileName) 
        return [s.name for s in doc.styles]
        
    

    
    def write_word(self, file_name:str,
                   text:str|None=None,
                   font_text:tuple[str,float,bool]=("Arial",11,False),
                   title:str|None=None,
                   heading:str="Heading 1",
                   font_title:tuple[str,float,bool]=("Arial",12,True),
                   write_list:list[str]|None=None,
                   table_maker:dict[str,list[Any]]|None=None)->bool:
        """
        DESCRIPTION: 
            This method will write new text in a word file.Will give an error message
            if the file does not exist. If a problem occurs, it will try to handle it

        PARAMETERS: 
            Moreover, you can add the following: 
            title: It is just a normal title 
            text: the main text the you want to write 
            table_maker: a normal table in word
            write_list: will create a basic enumerate list in the file

            font_text: basic configuration of the text
            font_title: basic configuration of the title
            heading: This is the Kind of head
        
        RETURN: 
            An String ('True') meaning it  wrote the file sucefully,
            otherwise doesn't

        EXAMPLE: 
            >>> FileManager().write_word('python.docx', 'Hello world!' )
                #In the python.docx file will appear Hello world!
                Output: 'True'
        
        
        """
        

        if not self.file_exist(file_name):
            print("why")
            print(f"World File {self.get_name(file_name)} don't found","You need to create a file first. "\
                              "Did you tried the FileManager().create_file() method?")
            return False
        
        if self.get_file_size(file_name) <= 0: 
            return False
        
        try:
            #reading the main file
            doc=Document(file_name)
            #
            font,size,bold=font_text
            

            #adding the heading
            if title is not None:
                hfont,hsize,hbold=font_title

                # Ensuring that Head is inside the word file, else a normal paragraph
                if heading in self.word_styles(file_name):
                    head = doc.add_heading(title)

                else: 

                    head = doc.add_paragraph(title)
                
                
                if head.runs: 
                    run = head.add_run(title)
                    run.font.name = hfont
                    run.font.size = Pt(hsize)
                    run.font.bold = hbold



            #Adding the paragraph
            if text is not None: 

                paragraph=doc.add_paragraph()
                run=paragraph.add_run(text + "\n")
                run.font.name=font
                run.font.size=Pt(size)
                run.font.bold=bold

            #Adding the numeral list
            if write_list is not None:

                for values in write_list:
                    list_=doc.add_paragraph()
                    run=list_.add_run(values,style="List Number")
                    run.font.name=font
                    run.font.size=Pt(size)
                    run.font.bold=bold

            if table_maker is not None: 
                keys:list[str]=list(table_maker.keys())
                content:list[list]=list(table_maker.values())

                

                table=doc.add_table(rows=len(content)+1,cols=len(keys),style="Table Grid")
                
      
                header=table.rows[0].cells
        
                for index,value in enumerate(keys): 
                    header[index].text = value

           

                for row_index, lts in enumerate(zip(*content),1):
                    data=table.rows[row_index].cells

                    for index, txt in enumerate(lts):
                        data[index].text=str(txt)

                    



            doc.save(file_name)
            
            
        
           

            self.storage_historial(file_name,"Write")
            print("Done", "Text was added to Word file. in the file: ",self.get_name(file_name))
            return True
 

        except Exception as e: 
            print("-"*90)
            print(f"Problem at: {e}: " + self.write_word.__name__)
            print("-"*90)

            if self._with_traceback:
                traceback.print_exc()
            return  False
    

           
 
    def create_word(self,fileName:str,text:str|None=None)->None: 
        """This will create a new file if the Does Not exist, it will show an error message
        otherwrise. Likewise you can write a text too, but this last is optional"""
        if self.file_exist(fileName): 
            print("File exist", f"'{self.get_name(fileName)}' already exist")
            return 


        if not fileName.endswith(".docx"): 
            fileName = fileName + ".docx"
        
    
        doc=Document()
            
        doc.save(fileName)
     
        if text is not None:
            self.write_word(fileName,text)
                


        self.storage_historial(fileName,"Create")

       
    def read_word(self, filePath)->str: 

        try: 
            with open(file=filePath,mode="r",encoding='utf-8') as file: 

                doc=Document(filePath)

                self.storage_historial(filePath,"Read")

                text= "\n".join(p.text for p in doc.paragraphs)

                tables= "" 

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tables + cell.text + "\t"
                        tables += "\n"
                


                return text + tables 
        
            


        except Exception as e: 
            print(e)

            return "N/A"
        



    #General methods 
    def create_txt(self,fileName:str,text:str|None=None)->bool: 
        """This will create a new file if the Does Not exist, it will show an error message
        otherwrise. Likewise you can write a text too, but this last is optional"""
        
        try: 
            with open(file=fileName, mode="x") as file:
                if text is not None: 
                    file.write(text)

            
                


            me=f"'{self.get_name(fileName)}' was created "
            print(me)
            message.showinfo("Complete",me)
            self.storage_historial(fileName,"Create")
            return True

                
                
        except FileExistsError as e: 
            print("File exist", f"'{self.get_name(fileName)}' already exist")
            return False
  
    def write_txt(self,fileName:str,text:str,mode_:Literal["a","w"]="a")->bool:
        """This will only add more text into the file. if mode_ is 'a', will not delete
         the previous info, will delete all the previous info if 'w' is provided """
  
        if not self.file_exist(fileName):
            print(f"File {self.get_name(fileName)} don't found","You need to create a file first. "\
                              "Did you tried the FileManager().create_file() method?")
            return False
        
        try: 
        
            with open(file=fileName, mode=mode_) as file:
                file.write(text)



            print("Done","The text was added in ", self.get_name(fileName))
            me=f"The text was added "
            print(me)
            self.storage_historial(fileName,"Write")
            return True

        except PermissionError as e : 
            print(f"{e}","Maybe you forgot to close the file?")
            return False

    def read_txt(self, filePath)->str: 

        try: 
            with open(file=filePath,mode="r",encoding='utf-8') as file: 
                self.storage_historial(filePath,"Read")
                return file.read()

        except Exception as e: 
            return "N/A"
        


    #json files
    def write_json(self,filePath:str,info:dict[str,Any])->bool:
        """This will only add more text into the json file.Appending the value 
        given to the last of the json file"""

        if not self.file_exist(filePath):
            print("File don't found","You need to create a file first. "\
                              "Did you tried the FileManager().create_file() method?")
            return False

        all_info=self.read_json(filePath,False)



        if  isinstance(all_info,list): 
            all_info.append(info)
        else: 
            all_info=[info]
            
        try: 
            with open(file=filePath,mode="w") as file:
                json.dump(all_info,file,indent=4)

                return True
    
                
                

        except Exception as e: 
            if self._with_traceback:
                traceback.print_exc()
            
            print("_"*90)
            print( "Problem at: " +  self.write_json.__name__)
            message.showerror("Error",f"{e}")
            print("_"*90)
            return False

    def create_json(self,filePath:str,text:dict[str,Any]|None=None)->bool: 
        """This will create a new json file if file Does Not exist, it will show an error message
        otherwrise. Likewise you can write a text too, but this last is optional"""
        if not filePath.endswith(".json"): 
            filePath = filePath + ".json"

        try: 
            with open(file=filePath,mode="x") as file:

                if text is not None: 
                    json.dump([text],file, indent=4)
                else: 
                    json.dump([],file,indent=4)

                print(f"'{self.get_name(filePath)}' was created ")
            
            self.storage_historial(filePath,"Create")

            return True


        except FileExistsError as e: 
            print("File exist", f"'{self.get_name(filePath)}' already exist")
            return False
        except Exception as e: 
            if self._with_traceback:
                traceback.print_exc()

            print("_"*90)
            print( "Problem at: " +  self.create_json.__name__)
            print("Error",f"{e}")
            print("_"*90)
            return False

    @overload
    def read_json(self, filePath:str,data_frame:Literal[False],/)->list[dict]: 
        ...

    @overload
    def read_json(self, filePath:str,data_frame:Literal[True],/)->pd.DataFrame: 
        ...


    def read_json(self, filePath:str,data_frame:bool=True): 
        """
        This methos will attempt to open a file in read mode('r') and return base of paramter given. 
        if a problem occurs, it will raturn an empty object based of the second parameter provided. 

        Return: 
            if data_frame is True, it will return a DataFrame object.
            if data_frame is False, it will return a list of dictionaries storage of the in the file
        """
        
        if data_frame: 
            data=pd.DataFrame()
        else: 
            data=[] #type:  ignore 

        try: 
            with open(file=filePath,mode="r",encoding='utf-8') as file: 
                #self.storage_historial(filePath,"Read")
                try: 
                    if data_frame:

                        return pd.read_json(file)
                    
                    else: 
                        return json.load(file)

                
                except Exception :
                    return data
              
        except Exception as e: 
            if self._with_traceback:
                traceback.print_exc()

            print( "Problem at: " +  self.read_json.__name__)
            print(e)
            #message.showerror("Error",f"{e}")
            return data

    def update_json(self,
                    filePath:str,
                    search_key:str,
                    search_value:Any,
                    **values_to_update)->None: 
        """
        DESCRIPTION:
            This method will read the file and attempt to update the value with the new value 
            provide.

        PARAMETERS:
            filePath: the main path where will atempt to search the file 
            search_key:i guess this wil be the main value where this method will try to 
            search. Since all the json objects will have the sames keys. so this is the only
            answer i could think 
            search_value: this is the value to find the specific dictionary. Since the majority 
            of time all the dicts will have the same key; Althoug some of value is very likely
             that have the same value. There is alwas a specific key that will diference them 

            values_to_update:to wrok more eficiency i think it would be better that work with **kwargs
            
            update_key: the spefific value you wanna update 
            new_value: the new value to be update
        
        """

        if not self.file_exist(filePath):
            print("The json file does not extist")

            return 


        datas=self.read_json(filePath,False)
        if not isinstance(datas,list):
            print("For the moment being i jsut can work with list, sorry")
            return 
        
        datas.sort(key=lambda b: b[search_key])
        index=my_dict_binary(datas,search_key,search_value)

        if index == -1: 
            print("The value was not found")
            return 
        
        try: 

            for key, value in values_to_update.items(): 
                datas[index][key] = value

         

            with open(filePath,mode="w") as file: 
                json.dump(datas,file,indent=4)


            print("The update was succesfully done!")




        except Exception as e: 
            print("Error: ",e)


    @property
    def withtraceback(self):
        return self._with_traceback

    @withtraceback.setter
    def withtraceback(self,value:bool):
        if not isinstance(value,bool): 
            raise ValueError("It must be a boolean")
        
        print(f"{self._with_traceback} change to {value}")
        self._with_traceback = value


"""
#22/08/2025
I wanna update a json file.
and i get structured: 
    1. First read the json file: 
        with open("my json.json","r") as file: 
            data=json.load(file)
    2. get the specific dict i want. The conclusion i reach was through i binary search :

        def my_dict_binary(dictionaries:list[dict[str,Any]],
                   key:str,
                   target:str)->int: 

            try: 
                start:int=0
                end:int=len(dictionaries)-1

                while start <= end:
                    mid:int= (start + end) // 2

                    middle_value = dictionaries[mid][key]
                    
                    if middle_value == target: return mid

                    if middle_value < target: start = mid +1

                    else: end = mid -1


                return -1

            except: 
                return -1

                
    data.sort(key=lambda d: d["Name"])
    index=my_dict_binary(data,"Name","Hersy") #2 i guess 

    3. Acces the specific key and 4. update the value
    data[index]["Country"] = "japan"

    5. save againthe json file: 
        with open("my json.json","w") as file: 
            json.dump(data,file,indent=4)


BUT I DUNNO IF THIS IS THE CORRECT APROACH. But please dont give the answer i wanna 
discover it by myself, just guide me please
        





"""







root=Tk()
root.withdraw()


        






file_path: str = r"C:\Users\Usuario Pc\Desktop\seminary budgets.xlsx"
my_text_txt: str = "C:\\Users\\Usuario Pc\\Downloads\\my text.txt"
python_docx:str = r"C:\Users\Usuario Pc\Desktop\python.docx"
python_xlsx:str = r"C:\Users\Usuario Pc\Desktop\python.xlsx"
file4: str = "Bye world!.docx"






if __name__ == '__main__': 
    

    """
    x = will crate a new file and we can write something too, however, if exist it'll raise an error
    w = will create a new file if it does not exist and we can write something too, 
    however, if it is exist it will delete all the privios info
    a= will NOT create a new file,only if it is exist we can write something,
    therefore, it will save all the previous info
    r= i dunno 
    
    """

    file=FileManager()
    texting:str="""
        ---

        ## *1. Theoretical Framework*

        ### *1.1. Project Title*

        Impact of Climate Change on the Right to Water in the Escuela Oficial Rural Mixta Río Azul de San Pedro Ayampuc.

        ---

        ### *1.2. Identification of the Problem*

        In Río Azul, *access to potable water is limited and irregular*. According to the document, even though some improvements were made, the water supply is not enough to cover basic needs.

        * In the school: classrooms and bathrooms cannot be cleaned properly, which threatens students’ health and spreads diseases.
        * Climate change worsens the problem: *droughts are longer and rainfall is more intense*, damaging infrastructure and reducing water sources.
        * Poor maintenance and limited investment make the situation worse.

        This undermines a *fundamental human right* recognized by the Guatemalan Constitution: the right to water.

        ---

        ### *1.3. Research Questions*

        Examples:

        * Why does scarcity persist and how does it affect the students?
        * Which human rights are violated without access to water?
        * What illnesses come from contaminated water?
        * How does climate change affect the school directly?
        * What policies could improve rural schools’ water management?

        ---

        ### *1.4. Formulation of the Problem*

        #### *1.4.1. Description*

        The lack of water affects hygiene, nutrition, and education. Children often cannot wash hands, clean classrooms, or stay hydrated. This reduces *academic performance and health*.

        #### *1.4.2. Contextualization*

        According to Interactive Country Fiches (2022), by 2050 in Guatemala:

        * Average temperature will rise between 0.75–2 °C.
        * Rainfall could decrease by 500–675 mm.
        * Water availability will drop by 26%.
        * 64% of the population will live under *water stress, and nearly half with **extreme scarcity*.

        This will deeply affect agriculture, poverty, and rural communities.

        #### *1.4.3. Reformulated Question*

        How can an awareness campaign at Río Azul School help students and the community protect the right to water against climate change?

        ---

        ### *1.5. Hypotheses of Action*

        * If at least *25% of students improve water habits*, teachers will be motivated to keep promoting water education.
        * Awareness campaigns may *reduce illnesses* (like gastroenteritis and diarrhea) by 90%.
        * If trained, *75% of students can apply home purification techniques* (filters, boiling).
        * Teaching students to care for bathrooms improves hygiene and reduces disease.
        * With recreational and educational activities, *85% of students will actively participate* in water conservation.

        ---
        ## *1.6. Literature Review*

        ### *a) Climate change produces water scarcity*

        * According to SGCCC (2019), climate change reduces both the *quantity and quality* of available water. It increases droughts, floods, and extreme weather events, which destabilize health, food security, and biodiversity.
        * Instituto del Agua highlights that deforestation, unsustainable farming, and urban growth worsen the crisis in Guatemala.

        ---

        ### *b) Decrease of water availability*

        * In Guatemala, scarcity has been a problem for more than a decade, especially in rural and poor areas.
        * According to Alice Burgor Paniagua, USAC (2016), regulations for wastewater and water protection exist but are *inefficient compared to other countries* like Mexico or Costa Rica.

        ---

        ### *c) Exhaustion of potable water*

        * Escobar (2023) notes that in some towns water arrives only a few times per week. Even in Guatemala City, service is irregular.
        * Despite the abundance of rivers and lakes, *6 of every 10 households lack potable water*.

        ---

        ### *d) Causes of water absence*

        * Main causes: population growth, poor water management, deforestation, climate change, and industries’ overuse.
        * According to Anderson (2020), Guatemala has abundant rainfall but distribution is unequal and mismanaged, creating shortages in highly populated areas such as Guatemala City.

        ---

        ### *e) Areas most affected by water scarcity*

        * According to the Consejo Internacional (2022), while Guatemala produces about *97 billion m³ of water annually, only **10% is effectively used*.
        * Urban areas: around *70% of households* have access to water services.
        * Rural areas: only *30% of households* access piped water and drainage.
        * This creates *inequality*: rural communities, especially women and children, are the ones who must fetch water daily.
        * Ve Viaja y Habla (2019) emphasizes that *95% of water in Guatemala is not safe for human consumption*. Children are most vulnerable: 1 in 20 rural children dies before age 5 due to waterborne diseases.

        ---

        ### *f) Consequences of scarcity*

        * Lack of water → diseases such as diarrhea, hepatitis A, cholera, and malnutrition.
        * Gómez (2009) states scarcity also reduces agricultural production, causing economic problems and even migration.
        * Ramírez (2024) adds that ecosystems (wetlands, rivers) disappear, leading to biodiversity loss.

        ---

        ### *g) Legal consequences*

        * According to the UN (2002, 2010), the right to water means access to sufficient, safe, acceptable, and affordable water.
        * However, USAC (2011) points out that Guatemala still lacks a *national Water Law*, which limits fair distribution and sustainable management.

        ---

        ### *h) Institutions responsible*

        * INFOM repairs water infrastructure.
        * The Ministry of Health seeks to prevent diseases from unsafe water.
        * Municipalities provide distribution, but often with *unequal coverage*.
        * USAC (2024) reports that *90% of rural water sources and 40% of urban ones do not meet quality standards*, exposing families to health risks.

        ---

        ### *i) Proposed solutions*

        * Stockholm Environment Institute (2024) suggests integrated watershed management, investment in infrastructure, and stronger policies.
        * Elizabeth (2024) highlights that only *2.5% of Earth’s water is freshwater*, so conservation and efficient use are urgent priorities.


        _
        ---

        ## *2. Methodological Design*

        ### *2.1. Type of Research*

        * The study used *Participatory Action Research (PAR)*.
        * This approach was chosen because it involves *direct work with the students* of Escuela Oficial Rural Mixta Río Azul.
        * Instead of only observing, the researchers encouraged students to participate in identifying problems and creating solutions.
        * PAR strengthens children’s capacities, gives them confidence, and creates lasting changes in the school environment.

        ---

        ### *2.2. Objectives*

        * *General Objective:*

        * Promote environmental awareness and responsibility in children so they understand the importance of protecting natural resources and mitigating climate change.

        * *Specific Objectives:*

        * Train children on the importance of natural resources and the impacts of climate change, to develop sustainable habits.
        * Carry out an *educational campaign* about climate change so students know how to act against it.
        * Teach children to *value and care for natural resources*.

        ---

        ### *2.3. Research Questions (Methodological Focus)*

        Examples of what the survey asked students:

        * How often does water arrive at your home?
        * What activity consumes the most water at home?
        * What do you think is the main cause of water pollution?
        * Which institutions are responsible for providing and protecting water?
        * Which natural resources are most affected by climate change?
        * Why has no definitive solution been found for water scarcity?
        * What diseases are most common due to contaminated water?
        * What is climate change?
        * Which human rights are violated by climate change?

        ---

        ### *2.4. Data Collection Method*

        * *Surveys* were used as the main tool.
        * Purpose: measure students’ knowledge and perceptions about climate change and water scarcity.
        * Allowed the researchers to obtain *quantitative and qualitative results* to guide educational activities.

        ---

        ### *2.5. Sample*

        * *Population:* students from Escuela Oficial Rural Mixta Río Azul (San Pedro Ayampuc, Guatemala).
        * *Target group:* primary-level children aged 10–13 years.
        * *Sample size:* 17 sixth-grade students were surveyed.

        ---

        ### *2.6. Survey (Results and Interpretation)*

        Some key findings:

        * *Water frequency at home:* Most students said it arrives every 2–5 days.
        * *Highest water consumption:* Washing clothes.
        * *Main cause of water contamination:* Solid waste and garbage.
        * *Responsible institution (in their view):* The municipality.
        * *Most affected resources by climate change:* Water, soil, and air.
        * *Cause of droughts:* Mainly climate change (11 students), second deforestation (6 students).
        * *Reasons for lack of solution to water scarcity:* Lack of citizen awareness (9 students) and lack of laws (5 students).
        * *Diseases from lack of water:* Dengue, diarrhea, vomiting.
        * *Definition of climate change (students’ view):* Increase in temperature.
        * *Rights violated:* Health and food rights.

        *Interpretation:*

        * Children are aware of water scarcity and its causes.
        * They associate poor management, waste, and climate change with the problem.
        * They recognize municipalities as responsible, but also note the absence of laws and citizen responsibility.

        ---

        ### *2.7. Implementation of Actions*

        * A *plan of action* was designed based on the findings.
        * Activities included:

        * Forming student groups.
        * Research on water and climate change.
        * Educational campaigns and talks.
        * Practical exercises on water purification and hygiene.
        * *Expected results:* stronger awareness, improved habits, and safer hygiene conditions at the school.

        ---

        ✅ *Summary of Methodological Design:*
        The project used *participatory action research, focusing on children as active participants. Through surveys, the team learned that water arrives irregularly (every 2–5 days), washing clothes consumes most water, and waste is seen as the main pollutant. Students identified **climate change and lack of awareness* as key reasons for scarcity. With this knowledge, the group implemented educational campaigns and practical activities to *promote responsible water use* and *protect health* in the school.




        ——

        ## *3. Implementation of the Action*

        ### *3.1. Planning and Preparation*

        The team carefully planned the intervention before applying it. They defined the objectives, organized student groups, and scheduled dates for activities. Preparation involved coordination with the school, the teacher in charge, and the seminar team to ensure the activities would be both educational and practical.

        ---

        ### *3.2. Intervention Strategies*

        Several strategies were selected to strengthen awareness among the students:

        * *Educational talks* explaining the impact of climate change and the importance of water.
        * *Practical demonstrations* on hygiene and purification methods.
        * *Interactive activities* such as games and group work to encourage participation.
        * *Visual aids* (posters, charts) to reinforce learning.

        ---

        ### *3.3. Characteristics of the Action*

        * *Scope:* The intervention targeted the sixth-grade students of Escuela Oficial Rural Mixta Río Azul, with the intention of extending the acquired knowledge to their families.
        * *Duration:* Activities were carried out across several weeks, combining research, campaign development, and awareness actions.
        * *Objective:* Encourage students to adopt responsible water-use habits and transmit this awareness to their community.
        * *Adaptation:* The strategies were adjusted to the children’s age, educational level, and the local reality of water scarcity.

        ---

        ### *3.4. Resources for Implementation*

        The project used *human resources* (students, teachers, researchers) and *material resources* such as computers, internet, printed material, markers, and filters for demonstrations. Time allocation was also considered as a fundamental resource.

        ---

        ### *3.5. Elements that Formed the Action*

        The action included:

        * Formation of student groups.
        * Research on climate change and water scarcity.
        * A structured campaign with talks, posters, and recreational activities.
        * Hygiene and cleaning practices within the school.
        * Demonstrations on water purification methods.

        ---

        ### *3.6. Description of the Action (Narrative)*

        The action began with group formation and the assignment of tasks. Students collaborated with the seminar team in researching information on climate change and water scarcity. Once this was completed, an *awareness campaign* was carried out inside the school.

        During the campaign, children participated in talks, created posters with slogans about water conservation, and received training on how to purify water using basic methods. In addition, they learned how to properly maintain the bathrooms and school facilities to avoid diseases. Recreational activities were also included to ensure motivation and engagement.

        The process allowed students not only to learn theoretical concepts but also to practice real actions that they could apply in their daily lives.

        ---

        ### *3.7. Checklist for Implementation*

        A checklist was designed to ensure that each stage of the action was completed. It included verification of planning, preparation of materials, execution of talks, development of practical activities, and evaluation of student participation.

        ---

        ✅ *Summary of Implementation of the Action:*
        The project was implemented through a *well-structured educational campaign* that combined research, awareness activities, and practical demonstrations. Students were actively involved in both learning and applying solutions. The intervention succeeded in linking *theory and practice*, helping students understand water scarcity not only as a concept but as a daily reality requiring responsible action.



        ---

        ## *4. Evaluation and Reflection*

        ### *4.1. Evaluation*

        The action was evaluated based on participation, student understanding, and changes in attitudes. The campaign was considered effective because students showed interest, collaborated in activities, and demonstrated greater awareness about responsible water use.

        ---

        ### *4.2. Reflection*

        * *Observation:* Students actively engaged in discussions, games, and demonstrations.
        * *Results Review:* They recognized the importance of water conservation and identified diseases caused by lack of hygiene.
        * *Analysis:* The campaign helped transform knowledge into practical actions (e.g., cleaning bathrooms, purifying water).
        * *Interpretation:* Children connected climate change with their local problem of water scarcity.
        * *Learning:* The process fostered responsibility and encouraged them to share knowledge with their families.

        ---

        ✅ *Crucial Idea:*
        The evaluation showed that *awareness campaigns are an effective tool* for teaching children about climate change and water scarcity. Students not only gained knowledge but also applied it in their school environment, creating a ripple effect toward their community.

        ---
        ——

        ---

        ## *5. Accountability*

        ### *5.1. Student Contributions*

        Each student in the seminar group contributed *Q41.00*, which was collected to finance the materials and activities necessary for the project.

        ---

        ### *5.2. Project Budget*

        The budget covered essential expenses such as:

        * Educational materials (markers, posters, paper).
        * Hygiene products (soap, disinfectant).
        * Printing of surveys and reports.
        * Transportation and miscellaneous costs.

        ---

        ### *5.3. Accountability Report*

        In addition to the students’ contributions, the project received *an additional income of Q200.00 from Librería Iris*, which helped cover extra expenses. All funds were carefully recorded, and their use was documented for transparency.

        ---

        ### *5.4. Closure (Finiquito)*

        At the end, a closure report (finiquito) was presented, certifying that the funds were managed responsibly and that the project was executed as planned.

        ---

        ✅ *Crucial Idea:*
        Through both student contributions and external support, the project ensured *responsible financial management*, teaching students about accountability and transparency in community projects.

        ---


        """
        

    #file.file_exist(python_docx)
    #file.file_exist("Hello world.txt")
    #file.create_file(file4, text="This is just a bye world!!")
    #file.write_file(file4,"Hello")
    #file.create_file(python_docx,".docx")

    #file.write_file(python_docx,"Python is cool!, you need to try it so bad !!")
    #file.file_exist(python_docx)

    #file.create_file("Hello json",".json",{"name":"Hersy","age":19})
    #file.write_file("Hello json.json",{"name":"Craice","age":18})

    #file.write_file(python_docx,text="Hello world! No.1 ")
    #file.write_file(python_docx,text="Hello world! NO.2  ")
    #file.write_file(python_docx,text="Hello world! No.3 .i guess so..")
    #file.write_file(python_docx,text="That's gotta be it..")
    text="""
    Things i lern: 
        x:
            will crate a new file and we can write something too, however, if exist it'll raise an error

        w:
            will create a new file if it does not exist and we can write something too, 
            however, if it is exist it will delete all the privios info

        a: 
            will NOT create a new file,only if it is exist we can write something,
            therefore, it will save all the previous info

        r:
            i dunno 
    
    """

    #file.write_file(python_docx,text)


    file.file_exist(python_docx)


    #file.write_file(file4,"Bye Nice world\n")


    #text=file.read_file(python_docx)
    #print(text)

    #file.write_file(python_docx,"What is wrong with me, and i dont even care about, if i where you ")

    #file.write_file(file4,texting)

    #file.write_file(file5,texting)
    
  




    """
    if file.file_exist(file4): 
        doc = Document(file4)

        tx = doc.paragraphs
        for i in tx: 
            print(i.text)

    """



 